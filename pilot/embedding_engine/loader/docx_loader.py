from typing import List, Optional

from langchain.docstore.document import Document
from langchain.document_loaders.base import BaseLoader
import docx


class DocxLoader(BaseLoader):
    """Load docx files."""

    def __init__(self, file_path: str, encoding: Optional[str] = None):
        """Initialize with file path."""
        self.file_path = file_path
        self.encoding = encoding

    def load(self) -> List[Document]:
        """Load from file path."""
        docs = []
        doc = docx.Document(self.file_path)
        content = []
        for i in range(len(doc.paragraphs)):
            para = doc.paragraphs[i]
            text = para.text
            content.append(text)
        docs.append(
            Document(page_content="".join(content), metadata={"source": self.file_path})
        )
        return docs
