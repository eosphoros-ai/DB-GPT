"""Docx Knowledge."""
from typing import Any, List, Optional

import docx

from dbgpt.core import Document
from dbgpt.rag.knowledge.base import (
    ChunkStrategy,
    DocumentType,
    Knowledge,
    KnowledgeType,
)


class DocxKnowledge(Knowledge):
    """Docx Knowledge."""

    def __init__(
        self,
        file_path: Optional[str] = None,
        knowledge_type: Any = KnowledgeType.DOCUMENT,
        encoding: Optional[str] = "utf-8",
        loader: Optional[Any] = None,
        **kwargs: Any,
    ) -> None:
        """Create Docx Knowledge with Knowledge arguments.

        Args:
            file_path(str,  optional): file path
            knowledge_type(KnowledgeType, optional): knowledge type
            encoding(str, optional): csv encoding
            loader(Any, optional): loader
        """
        self._path = file_path
        self._type = knowledge_type
        self._loader = loader
        self._encoding = encoding

    def _load(self) -> List[Document]:
        """Load docx document from loader."""
        if self._loader:
            documents = self._loader.load()
        else:
            docs = []
            doc = docx.Document(self._path)
            content = []
            for i in range(len(doc.paragraphs)):
                para = doc.paragraphs[i]
                text = para.text
                content.append(text)
            docs.append(
                Document(content="\n".join(content), metadata={"source": self._path})
            )
            return docs
        return [Document.langchain2doc(lc_document) for lc_document in documents]

    @classmethod
    def support_chunk_strategy(cls) -> List[ChunkStrategy]:
        """Return support chunk strategy."""
        return [
            ChunkStrategy.CHUNK_BY_SIZE,
            ChunkStrategy.CHUNK_BY_PARAGRAPH,
            ChunkStrategy.CHUNK_BY_SEPARATOR,
        ]

    @classmethod
    def default_chunk_strategy(cls) -> ChunkStrategy:
        """Return default chunk strategy."""
        return ChunkStrategy.CHUNK_BY_SIZE

    @classmethod
    def type(cls) -> KnowledgeType:
        """Return knowledge type."""
        return KnowledgeType.DOCUMENT

    @classmethod
    def document_type(cls) -> DocumentType:
        """Return document type."""
        return DocumentType.DOCX
