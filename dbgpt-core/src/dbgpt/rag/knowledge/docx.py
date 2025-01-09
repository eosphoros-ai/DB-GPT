"""Docx Knowledge."""
from typing import Any, Dict, List, Optional, Union

import docx
from docx.opc.oxml import parse_xml
from docx.opc.pkgreader import _SerializedRelationship, _SerializedRelationships

from dbgpt.core import Document
from dbgpt.rag.knowledge.base import (
    ChunkStrategy,
    DocumentType,
    Knowledge,
    KnowledgeType,
)


def load_from_xml_v2(base_uri, rels_item_xml):
    """Return |_SerializedRelationships| instance loaded with the relationships.

    contained in *rels_item_xml*.collection if *rels_item_xml* is |None|.
    """
    srels = _SerializedRelationships()
    if rels_item_xml is not None:
        rels_elm = parse_xml(rels_item_xml)
        for rel_elm in rels_elm.Relationship_lst:
            if rel_elm.target_ref in ("../NULL", "NULL"):
                continue
            srels._srels.append(_SerializedRelationship(base_uri, rel_elm))
    return srels


class DocxKnowledge(Knowledge):
    """Docx Knowledge."""

    def __init__(
        self,
        file_path: Optional[str] = None,
        knowledge_type: Any = KnowledgeType.DOCUMENT,
        encoding: Optional[str] = "utf-8",
        loader: Optional[Any] = None,
        metadata: Optional[Dict[str, Union[str, List[str]]]] = None,
        **kwargs: Any,
    ) -> None:
        """Create Docx Knowledge with Knowledge arguments.

        Args:
            file_path(str,  optional): file path
            knowledge_type(KnowledgeType, optional): knowledge type
            encoding(str, optional): csv encoding
            loader(Any, optional): loader
        """
        super().__init__(
            path=file_path,
            knowledge_type=knowledge_type,
            data_loader=loader,
            metadata=metadata,
            **kwargs,
        )
        self._encoding = encoding

    def _load(self) -> List[Document]:
        """Load docx document from loader."""
        if self._loader:
            documents = self._loader.load()
        else:
            docs = []
            _SerializedRelationships.load_from_xml = load_from_xml_v2  # type: ignore
            doc = docx.Document(self._path)
            content = []

            for i in range(len(doc.paragraphs)):
                para = doc.paragraphs[i]
                text = para.text
                content.append(text)
            metadata = {"source": self._path}
            if self._metadata:
                metadata.update(self._metadata)  # type: ignore
            docs.append(Document(content="\n".join(content), metadata=metadata))
            return docs
        return [Document.langchain2doc(lc_document) for lc_document in documents]

    @classmethod
    def support_chunk_strategy(cls) -> List[ChunkStrategy]:
        """Return support chunk strategy."""
        return [
            ChunkStrategy.CHUNK_BY_SIZE,
            ChunkStrategy.CHUNK_BY_PARAGRAPH,
            ChunkStrategy.CHUNK_BY_SEPARATOR,
        ]

    @classmethod
    def default_chunk_strategy(cls) -> ChunkStrategy:
        """Return default chunk strategy."""
        return ChunkStrategy.CHUNK_BY_SIZE

    @classmethod
    def type(cls) -> KnowledgeType:
        """Return knowledge type."""
        return KnowledgeType.DOCUMENT

    @classmethod
    def document_type(cls) -> DocumentType:
        """Return document type."""
        return DocumentType.DOCX
